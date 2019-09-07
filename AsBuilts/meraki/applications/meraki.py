#!/usr/bin/python3
import json
import requests
from requests import urllib3
import time
import pprint as pp
import csv
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ..models import MerakiInfo
import os
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)


def http_get(meraki_url):
    base_url = ('https://api.meraki.com/api/v0/{}'.format(meraki_url))
    for api in MerakiInfo.objects.all():
        api_key = api.api_key
    headers = {'X-Cisco-Meraki-API-Key': api_key}
    try:
        get_response = requests.get(base_url, headers=headers, verify=False)
        status = get_response.status_code
        get_response = get_response.json()
        time.sleep(0.5)
    except:
        print('Meraki Cloud not reachable - check connection')
        get_response = 'unreachable'
        status = 'unreachable'

    return get_response, status


def create_word_doc_title():
    doc = Document('media/Meraki As Built.docx')
    doc.add_page_break()

    return doc

def ins_word_doc_image(doc, pic_dir, pic_width=5.25):
    doc.add_picture(pic_dir, width=Inches(pic_width))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    return doc


def create_word_doc_paragraph(doc, heading_text = '', heading_level = 1,
                            paragraph_text = ''):
    doc.add_heading(heading_text, level=heading_level)
    p = doc.add_paragraph(paragraph_text)

    return doc

def create_word_doc_text(paragraph_text, doc):
    p = doc.add_paragraph(paragraph_text)

    return doc

def create_word_doc_bullet(
    doc,
    bp1 = '',
    bp2 = '',
    bp3 = '',
    bp4 = '',
    bp5 = ''
    ):
    if bp1 != '':
        doc.add_paragraph(
            bp1, style='List Paragraph'
        )
    if bp2 != '':
        doc.add_paragraph(
            bp2, style='List Paragraph'
        )
    if bp3 != '':
        doc.add_paragraph(
            bp3, style='List Paragraph'
        )
    if bp4 != '':
        doc.add_paragraph(
            bp4, style='List Paragraph'
        )
    if bp5 != '':
        doc.add_paragraph(
            bp5, style='List Paragraph'
        )


    return doc


def create_word_doc_table(doc, df):
    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
    #if isinstance(df, pd.DataFrame) and df.empty == False:
    print(df)
    try:
        t = doc.add_table(df.shape[0]+1, df.shape[1], style = 'Grid Table 4 Accent 6')
        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]
        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
    except:
        doc = doc
        print('Unable to add table')

    doc.add_page_break()
    return doc


def save_word_document(doc, customer):
    doc.save('media/tmp/{}-AS_Built.docx'.format(customer))

def get_network_info(input_dict, append_url = '', dict_key ='', list=True):
    for network in input_dict:
        try:
            data, status_code = http_get(meraki_url='networks/{}/{}'.format(network["id"],
            append_url))
            if list == True:
                for i in data:
                    if 'errors' not in i:
                        print(input_dict[dict_key])
                        input_dict[dict_key].append(i)
                    else:
                        print("error found")
            elif list == False:
                input_dict[dict_key].append(data)
        except:
            print('{} not reachable'.format(append_url))


def pull_data(meraki_url='organizations'):
    # Intitialise dictionary to store returned data
    data, status_code = http_get(meraki_url)
    return data, status_code


def get_org_info(dn='networks'):
    data = []
    orgs, status_code = pull_data(meraki_url='organizations')
    try:
        data_df = pd.DataFrame.from_dict(orgs)
    except:
        data_df = pd.DataFrame.from_dict(orgs, orient='index')
    for org in orgs:
        # Get networks from all orgs and add to dictionary
        org_data, status_code = pull_data(
            meraki_url='organizations/{}/{}'.format(org["id"], dn))
        log = ('organizations/{}/{}  - Returned status code: {}  '.format(
            org["id"], dn, status_code))
        print(log)
        for i in org_data:
            data.append(i)
        if len(org_data) > 0:
            try:
                out_df = pd.DataFrame.from_dict(org_data)
            except:
                out_df = pd.DataFrame.from_dict(org_data, orient='index')
    return data, log, data_df, out_df


def get_network_info(network, append_url = ''):
    data, status_code = http_get(meraki_url='networks/{}/{}'.format(network["id"],
    append_url))
    log = ('networks/{}/{}  - Returned status code: {} '.format(network["id"],
    append_url, status_code))
    print(log)
    #if isinstance(data, dict):
    try:
        data_df = pd.DataFrame.from_dict(data)
    except:
        try:
            data_df = pd.DataFrame.from_dict(data, orient='index')
        except:
            data_df = data
    return data_df, log
