from docx import Document
from docx.shared import Inches

doc = Document()

doc.add_heading('Document Title', 0)

p = doc.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

doc.add_heading('Heading, level 1', level=1)
doc.add_paragraph('Intense quote', style='Intense Quote')

doc.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
doc.add_paragraph(
    'first item in ordered list', style='List Number'
)

doc.add_picture('monty-truth.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

# add a table to the end and create a reference variable
# extra row is so we can add the header row
t = doc.add_table(df.shape[0]+1, df.shape[1])

# add the header rows.
for j in range(df.shape[-1]):
    t.cell(0,j).text = df.columns[j]

# add the rest of the data frame
for i in range(df.shape[0]):
    for j in range(df.shape[-1]):
        t.cell(i+1,j).text = str(df.values[i,j])

doc.add_page_break()

doc.save('demo.docx')
