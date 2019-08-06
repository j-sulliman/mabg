#!/usr/bin/python3
import json
import requests
from requests import urllib3
import time
import pprint as pp
import csv
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT, WD_SECTION
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)


def http_get(meraki_url):
    base_url = ('https://api.meraki.com/api/v0/{}'.format(meraki_url))
    headers = {'X-Cisco-Meraki-API-Key': '029c931f51ee7ea3d74de7cbf6e2db80b17d38d7'}
    get_response = requests.get(base_url, headers=headers, verify=False).json()
    time.sleep(0.5)
    return get_response


def create_word_doc_title(doc_title = ''):
    doc = Document()
    doc.add_heading(doc_title, 0)
    doc.add_picture('meraki_splash.png', width=Inches(3.25))
    doc.add_page_break()

    return doc


def create_word_doc_paragraph(doc, heading_text = '', heading_level = 1,
                            paragraph_text = ''):
    doc.add_heading(heading_text, level=heading_level)
    #new_section.orientation = WD_ORIENT.LANDSCAPE
    p = doc.add_paragraph(paragraph_text)
    current_section = doc.sections[-1]
    new_section = doc.add_section()
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = new_section.page_height, new_section.page_width
    new_section.page_width = new_width
    new_section.page_height = new_height

    return doc


def create_word_doc_table(doc, df):
    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
    t = doc.add_table(df.shape[0]+1, df.shape[1], style = 'Light List Accent 1')

    # add the header rows.
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]



    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])

    doc.add_page_break()
    return doc



def main():
    # Intitialise dictionary to store returned data
    meraki_data = {}
    orgs = http_get(meraki_url='organizations')
    meraki_data['orgs'] = orgs
    meraki_data["networks"] = []
    meraki_data["licenses"] = []
    meraki_data["devices"] = []
    meraki_data["securityEvents"] = []
    meraki_data["ssids"] = []
    meraki_data["fwServices"] = []
    meraki_data["intrusionPrevention"] = []
    meraki_data["malwareSettings"] = []
    meraki_data["l3FirewallRules"] = []
    meraki_data["vlans"] = []
    meraki_data["traffic"] = []
    meraki_data["clients"] = []


    for org in meraki_data['orgs']:
        # Get networks from all orgs and add to dictionary
        networks = http_get(meraki_url='organizations/{}/networks'.format(org["id"]))
        for network in networks:
            meraki_data["networks"].append(network)

        # Get license info from all orgs and add to dictionary
        licenses = http_get(meraki_url='organizations/{}/licenseState'.format(org["id"]))
        meraki_data["licenses"].append(licenses)

        # Get IPs
        ips = http_get(meraki_url='organizations/{}/security/intrusionSettings'.format(org["id"]))
        for ip in ips:
            meraki_data["intrusionPrevention"].append(ip)

    for network in meraki_data['networks']:
        def get_network_info(append_url = '', dict_key =''):
            data = http_get(meraki_url='networks/{}/{}'.format(network["id"],
            append_url))
            for i in data:
                meraki_data[dict_key].append(i)
    get_network_info(append_url='devices', dict_key='devices')
    get_network_info(append_url='securityEvents', dict_key='securityEvents')
    get_network_info(append_url='ssids', dict_key='ssids')
    get_network_info(append_url='firewalledServices', dict_key='fwServices')
    get_network_info(append_url='l3FirewallRules', dict_key='l3FirewallRules')
    get_network_info(append_url='vlans', dict_key='vlans')
    get_network_info(append_url='traffic', dict_key='traffic')
    get_network_info(append_url='clients', dict_key='clients')


    licenses_df = pd.DataFrame.from_dict(meraki_data['licenses'])
    orgs_df = pd.DataFrame.from_dict(meraki_data['orgs'])
    networks_df = pd.DataFrame.from_dict(meraki_data['networks'])
    meraki_vlans_df = pd.DataFrame.from_dict(meraki_data['vlans'])
    l3FirewallRules_df = pd.DataFrame.from_dict(meraki_data['l3FirewallRules'])
    ssids_df = pd.DataFrame.from_dict(meraki_data['ssids'])
    devices_df = pd.DataFrame.from_dict(meraki_data['devices'])
    clients_df = pd.DataFrame.from_dict(meraki_data['clients'])
    traffic_df = pd.DataFrame.from_dict(meraki_data['traffic'])
    pp.pprint (meraki_data)

    with pd.ExcelWriter('output.xlsx') as writer:  # doctest: +SKIP
        licenses_df.to_excel(writer, sheet_name='licenses')
        meraki_vlans_df.to_excel(writer, sheet_name='VLANs')
        l3FirewallRules_df.to_excel(writer, sheet_name='l3FirewallRules')
        devices_df.to_excel(writer, sheet_name='devices')
        #l3FirewallRules_df.to_excel(writer, sheet_name='l3FirewallRules')
        ssids_df.to_excel(writer, sheet_name='ssids')
        clients_df.to_excel(writer, sheet_name='clients')
        traffic_df.to_excel(writer, sheet_name='traffic')

    doc = create_word_doc_title(doc_title = 'Meraki AS Built')
    doc = create_word_doc_paragraph(doc = doc, heading_text = 'License Overview',
                            paragraph_text='The following licenses were found:')
    doc = create_word_doc_table(doc, licenses_df)
    doc = create_word_doc_paragraph(doc = doc, heading_text = 'Device Overview',
                            paragraph_text='The following devices were found:')
    doc = create_word_doc_table(doc, devices_df)
    doc = create_word_doc_paragraph(doc = doc, heading_text = 'Layer-2',
                            paragraph_text='The following VLANs are configured:')
    doc = create_word_doc_table(doc, meraki_vlans_df)
    doc = create_word_doc_paragraph(doc = doc, heading_text = 'Firewall Rules',
                            paragraph_text='The following Rules are configured:')
    doc = create_word_doc_table(doc, l3FirewallRules_df)
    doc = create_word_doc_paragraph(doc = doc, heading_text = 'Wireless',
                            paragraph_text='The following SSIDs are configured:')
    doc = create_word_doc_table(doc, ssids_df)
    doc = create_word_doc_paragraph(doc = doc, heading_text = 'Client Overview',
                            paragraph_text='The following Clients were found:')
    doc = create_word_doc_table(doc, clients_df)
    '''
    doc = create_word_doc_paragraph(doc = doc, heading_text = 'Traffic Overview',
                            paragraph_text='Below shows an overview of traffic on the network:')
    doc = create_word_doc_table(doc, traffic_df)
    '''
    doc.save('demo.docx')

main()
